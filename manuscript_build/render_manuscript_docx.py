"""Render the manuscript to .docx in the same design as the PDF (DESIGN_SPEC.md).

This replaces `manuscript_docx/build_manuscript_docx.py` for the current manuscript. That builder
implements a different design system — a memo-style masthead, a hardcoded title block, a draft-status
note — and it silently dropped tables that were not introduced by a `### Table N.` caption. It is kept
only for the superseded builds it produced.

Font caveat: Word does not embed fonts written by python-docx, so Charis SIL renders as Charis SIL only
on a machine that has it installed. Readers without it get their Word default substitute and the page
will reflow. The PDF is the fidelity-correct artifact; the .docx is for people who want to comment or
track changes.

Usage:
    python render_manuscript_docx.py SOURCE.md OUTPUT.docx
"""
from __future__ import annotations

import pathlib
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ---- design constants, shared with render_manuscript_pdf.py -------------------------------------
SERIF = "Charis SIL"
MONO = "Courier New"
BODY_SIZE = 10.0
BODY_LEADING = 14.2
PARA_SPACE = 10.0
CODE_SIZE = 8.5
TABLE_SIZE = 8.0
HEADING_SIZES = {1: 15.0, 2: 12.5, 3: 11.0, 4: 10.5}
MARGIN_IN = 64.0 / 72.0        # 64 pt
BLACK = RGBColor(0, 0, 0)

_CODE = re.compile(r"`([^`]+)`")
_BOLD = re.compile(r"\*\*(.+?)\*\*", re.S)
_ITAL = re.compile(r"(?<![\*\w])\*(?!\s)([^*]+?)(?<!\s)\*(?!\*)", re.S)
_SUP = re.compile(r"\^([^\^\s]+)\^")
_LINK = re.compile(r"\[([^\]]+)\]\([^)]*\)")
_TOKEN = re.compile(r"(\x00\d+\x00|\*\*.+?\*\*|(?<![\*\w])\*[^*]+?\*(?!\*)|\^[^\^\s]+\^)", re.S)


def add_runs(par, text: str, size: float = BODY_SIZE, bold_all: bool = False,
             italic_all: bool = False) -> None:
    """Write Markdown inline markup into a paragraph as formatted runs."""
    stash: list[str] = []

    def keep(m: re.Match) -> str:
        stash.append(m.group(1))
        return f"\x00{len(stash) - 1}\x00"

    text = _CODE.sub(keep, text)
    text = _LINK.sub(r"\1", text)

    for piece in _TOKEN.split(text):
        if not piece:
            continue
        bold, italic, sup, mono = bold_all, italic_all, False, False
        m = re.fullmatch(r"\x00(\d+)\x00", piece)
        if m:
            piece, mono = stash[int(m.group(1))], True
        elif piece.startswith("**") and piece.endswith("**"):
            piece, bold = piece[2:-2], True
        elif piece.startswith("^") and piece.endswith("^"):
            piece, sup = piece[1:-1], True
        elif piece.startswith("*") and piece.endswith("*"):
            piece, italic = piece[1:-1], True
        run = par.add_run(piece)
        run.font.name = MONO if mono else SERIF
        run.font.size = Pt(CODE_SIZE if mono else size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = BLACK
        if sup:
            run.font.superscript = True
        # python-docx sets only the latin face; set the complex/east-asian faces too so Word does
        # not fall back to Calibri for any run.
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), MONO if mono else SERIF)


DOUBLE_SPACED = False


def style_paragraph(par, space_after: float = PARA_SPACE) -> None:
    pf = par.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    if DOUBLE_SPACED:
        pf.line_spacing = 2.0
    else:
        pf.line_spacing = Pt(BODY_LEADING)
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT


def repeat_header_row(row) -> None:
    """Mark a table's first row as a header so Word repeats it on every continuation page.

    Long tables legitimately span pages at PLOS, but a continuation fragment with no header is
    unreadable. This sets w:tblHeader, which the earlier builds did not.
    """
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def set_cell_border_bottom(cell, size: int = 6) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    tcpr.append(borders)


def split_row(line: str) -> list[str]:
    parts, cur, i = [], [], 0
    while i < len(line):
        c = line[i]
        if c == "\\" and i + 1 < len(line) and line[i + 1] == "|":
            cur.append("|")
            i += 2
            continue
        if c == "|":
            parts.append("".join(cur))
            cur = []
            i += 1
            continue
        cur.append(c)
        i += 1
    parts.append("".join(cur))
    if parts and not parts[0].strip():
        parts = parts[1:]
    if parts and not parts[-1].strip():
        parts = parts[:-1]
    return [p.strip() for p in parts]


def is_separator(cells: list[str]) -> bool:
    return bool(cells) and all(re.fullmatch(r":?-{2,}:?", c.replace(" ", "")) for c in cells)


def add_line_numbers(section) -> None:
    """Continuous line numbers that do not restart per page, which PLOS requires for review."""
    sectPr = section._sectPr
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:restart"), "continuous")
    ln.set(qn("w:distance"), "360")
    sectPr.append(ln)


def add_page_numbers(section) -> None:
    par = section.footer.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run()
    for el, attrs, text in (("w:fldChar", {"w:fldCharType": "begin"}, None),
                            ("w:instrText", {"xml:space": "preserve"}, " PAGE "),
                            ("w:fldChar", {"w:fldCharType": "end"}, None)):
        e = OxmlElement(el)
        for k, v in attrs.items():
            e.set(qn(k), v)
        if text:
            e.text = text
        run._element.append(e)
    run.font.name = SERIF
    run.font.size = Pt(BODY_SIZE)


def build(source: Path, output: Path, submission: bool = False) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)      # US Letter
    sec.left_margin = sec.right_margin = Inches(MARGIN_IN)
    sec.top_margin = sec.bottom_margin = Inches(MARGIN_IN)

    normal = doc.styles["Normal"]
    normal.font.name = SERIF
    normal.font.size = Pt(BODY_SIZE)
    normal.font.color.rgb = BLACK

    if submission:
        add_line_numbers(sec)
        add_page_numbers(sec)

    text = source.read_text()
    if submission:
        # PLOS uploads figures as separate files; the manuscript carries captions only.
        text = re.sub(r"^!\[[^\]]*\]\([^)]*\)\n\n?", "", text, flags=re.M)
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue

        m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if m:
            path = (source.parent / m.group(2)).resolve()
            if path.exists():
                doc.add_picture(str(path), width=Inches(8.5 - 2 * MARGIN_IN))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue

        m = re.match(r"^(#{1,4}) +(.*)$", line)
        if m:
            level = len(m.group(1))
            par = doc.add_paragraph()
            style_paragraph(par, space_after=6.0)
            par.paragraph_format.space_before = Pt(0 if level == 1 else 12.0)
            add_runs(par, m.group(2), size=HEADING_SIZES[level],
                     bold_all=True, italic_all=(level == 4))
            i += 1
            continue

        if line.lstrip().startswith("|"):
            block = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                block.append(lines[i])
                i += 1
            rows = [split_row(b) for b in block]
            rows = [r for r in rows if not is_separator(r)]
            if not rows:
                continue
            ncol = max(len(r) for r in rows)
            rows = [r + [""] * (ncol - len(r)) for r in rows]
            table = doc.add_table(rows=len(rows), cols=ncol)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = True
            for ri, row in enumerate(rows):
                for ci, cell_text in enumerate(row):
                    cell = table.cell(ri, ci)
                    par = cell.paragraphs[0]
                    style_paragraph(par, space_after=2.0)
                    par.paragraph_format.line_spacing = Pt(10.4)
                    add_runs(par, cell_text, size=TABLE_SIZE, bold_all=(ri == 0))
                    if ri == 0:
                        set_cell_border_bottom(cell)
            if table.rows:
                repeat_header_row(table.rows[0])
            style_paragraph(doc.add_paragraph(), space_after=PARA_SPACE)
            continue

        m = re.match(r"^(\s*)([-*+]|\d+\.) +(.*)$", line)
        if m:
            ordered = m.group(2) not in ("-", "*", "+")
            while i < len(lines):
                mm = re.match(r"^(\s*)([-*+]|\d+\.) +(.*)$", lines[i])
                if not mm:
                    break
                par = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
                style_paragraph(par, space_after=3.0)
                add_runs(par, mm.group(3))
                i += 1
            continue

        if re.fullmatch(r"(-{3,}|\*{3,}|_{3,})", line.strip()):
            i += 1
            continue

        para = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            if (not nxt.strip() or nxt.lstrip().startswith("|")
                    or re.match(r"^#{1,6} ", nxt) or nxt.startswith("![")
                    or re.match(r"^(\s*)([-*+]|\d+\.) +", nxt)):
                break
            para.append(nxt)
            i += 1
        par = doc.add_paragraph()
        style_paragraph(par)
        add_runs(par, " ".join(x.strip() for x in para))

    doc.core_properties.title = next(
        (l[2:].strip() for l in lines if l.startswith("# ")), "")
    doc.core_properties.author = ", ".join(
        re.findall(r"\*\*([^*]+)\*\*",
                   next((l for l in lines if l.startswith("**") and "^1^" in l), "")))
    doc.core_properties.subject = ("Exploratory secondary analysis of two published phosphosite-mutant "
                                   "screens, in yeast and human")
    doc.core_properties.comments = ""
    # python-docx's default template stamps 2013-12-23 and leaves keywords blank. Both ship to the
    # journal inside the file's properties, so they are set from the Markdown rather than left as
    # template residue. The timestamp is the source's own modification time, so a rebuild of an
    # unchanged source produces an unchanged property.
    kw = next((ln[len("**Keywords:**"):].replace(";", ",").strip()
               for ln in text.splitlines() if ln.startswith("**Keywords:**")), "")
    if kw:
        doc.core_properties.keywords = kw
    import datetime
    stamp = datetime.datetime.fromtimestamp(
        Path(source).stat().st_mtime, datetime.timezone.utc).replace(tzinfo=None)
    doc.core_properties.created = stamp
    doc.core_properties.modified = stamp
    doc.core_properties.last_modified_by = doc.core_properties.author
    doc.save(str(output))


if __name__ == "__main__":
    argv = [a for a in sys.argv[1:] if a != "--submission"]
    if len(argv) != 2:
        raise SystemExit(__doc__)
    sub = "--submission" in sys.argv
    if sub:
        DOUBLE_SPACED = True
        globals()["DOUBLE_SPACED"] = True
    build(Path(argv[0]).resolve(), Path(argv[1]).resolve(), submission=sub)
    print(f"wrote {argv[1]}" + (" (submission format)" if sub else ""))
